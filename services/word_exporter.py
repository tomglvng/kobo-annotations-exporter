from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sanitize_filename import sanitize
from pathlib import Path

from services.exporter_interface import ExporterInterface


class WordExporter(ExporterInterface):
    def export(self, retrieved_annotations: dict, directory: str) -> None:
        for author in retrieved_annotations:
            author_directory = "{}/{}/".format(directory, sanitize(author))
            Path(author_directory).mkdir(parents=True, exist_ok=True)
            books = retrieved_annotations[author]
            for book in books:
                book_file_name = "{}/{}.docx".format(author_directory, sanitize(book))
                document = Document()
                document.add_heading(book, level=1)
                document.add_paragraph(author, style='Caption', )
                p_blank = document.add_paragraph("")
                p_blank.line_spacing_rule = WD_LINE_SPACING.DOUBLE

                chapters = books[book]
                for chapter in chapters:
                    document.add_paragraph(chapter, style='Title')
                    p_blank = document.add_paragraph("")
                    p_blank.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    annotations = chapters[chapter]
                    for annotation in annotations:
                        comment = ''
                        if annotation.comment is not None and annotation.comment:
                            comment = annotation.comment
                        p_annotation = document.add_paragraph(annotation.text, style='Intense Quote')
                        p_annotation.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        p_comment = document.add_paragraph(comment, style='No Spacing')
                        p_comment.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        p_blank = document.add_paragraph("")
                        p_blank.line_spacing_rule = WD_LINE_SPACING.DOUBLE

                document.save(book_file_name)
